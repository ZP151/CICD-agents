from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).parent
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6673"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
WHITE = "FFFFFF"
PAGE_WIDTH = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for edge in ("top", "start", "bottom", "end"):
                node = tc_mar.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "80" if edge in ("top", "bottom") else "120")
                node.set(qn("w:type"), "dxa")


def set_font(run, size=11, color="000000", bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(paragraph, before=0, after=6, line=1.25, alignment=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def add_text(doc, text, *, bold_prefix=None, after=6, italic=False, color="000000"):
    p = doc.add_paragraph()
    style_paragraph(p, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_font(run, bold=True, color=color)
        run = p.add_run(text[len(bold_prefix):])
        set_font(run, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_font(run, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        size, color, before, after = 16, BLUE, 18, 10
    elif level == 2:
        size, color, before, after = 13, BLUE, 14, 7
    else:
        size, color, before, after = 12, DARK_BLUE, 10, 5
    style_paragraph(p, before=before, after=after)
    set_font(p.add_run(text), size=size, color=color, bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        fmt = p.paragraph_format
        fmt.left_indent = Inches(0.375)
        fmt.first_line_indent = Inches(-0.188)
        fmt.space_after = Pt(4)
        fmt.line_spacing = 1.25
        set_font(p.add_run(item))


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        fmt = p.paragraph_format
        fmt.left_indent = Inches(0.375)
        fmt.first_line_indent = Inches(-0.188)
        fmt.space_after = Pt(4)
        fmt.line_spacing = 1.25
        set_font(p.add_run(item))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.1)
        set_font(p.add_run(header), size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.1)
            set_font(p.add_run(value), size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [PAGE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GRAY)
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.2)
    set_font(p.add_run(f"{label}  "), size=10.5, color=INK, bold=True)
    set_font(p.add_run(body), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc, prompt):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [PAGE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FB")
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.15)
    run = p.add_run(prompt)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(9.2)
    run.font.color.rgb = RGBColor.from_string("20324A")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def setup_doc(header_text, footer_text):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    header = section.header.paragraphs[0]
    style_paragraph(header, after=0, line=1.0)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run(header_text), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    style_paragraph(footer, after=0, line=1.0)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run(footer_text), size=8.5, color=MUTED)
    return doc


def add_title(doc, title, subtitle, stripe):
    p = doc.add_paragraph()
    style_paragraph(p, before=6, after=4)
    set_font(p.add_run(title), size=25, color=INK, bold=True)
    p = doc.add_paragraph()
    style_paragraph(p, after=16)
    set_font(p.add_run(subtitle), size=13, color=MUTED)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [3120, 3120, 3120])
    for cell, (label, value) in zip(table.rows[0].cells, stripe):
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.1)
        set_font(p.add_run(label + "\n"), size=8.5, color=MUTED, bold=True)
        set_font(p.add_run(value), size=10, color=INK, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_architecture(doc, labels):
    add_heading(doc, labels["architecture"], 1)
    add_callout(doc, labels["message"], labels["architecture_message"])
    rows = [
        (labels["desktop"], labels["desktop_desc"]),
        (labels["daemon"], labels["daemon_desc"]),
        (labels["runtime"], labels["runtime_desc"]),
        (labels["integrations"], labels["integrations_desc"]),
        (labels["policy"], labels["policy_desc"]),
    ]
    add_table(doc, [labels["layer"], labels["responsibility"]], rows, [2250, 7110])
    add_text(doc, labels["flow"], bold_prefix=labels["flow_label"], after=8)


def add_turn(doc, labels):
    add_heading(doc, labels["turn"], 1)
    add_callout(doc, labels["speaker_line"], labels["turn_quote"])
    add_numbers(doc, labels["turn_steps"])
    add_text(doc, labels["turn_guard"], bold_prefix=labels["guard_label"], after=8)


def add_agenda(doc, labels):
    add_heading(doc, labels["agenda"], 1)
    add_table(doc, [labels["time"], labels["segment"], labels["talk_track"]], labels["agenda_rows"], [900, 2450, 6010])
    add_callout(doc, labels["rule"], labels["rule_text"])


def add_demo(doc, labels):
    add_heading(doc, labels["demo"], 1)
    for item in labels["demos"]:
        add_heading(doc, item["title"], 2)
        add_text(doc, item["goal"], bold_prefix=labels["goal_label"], after=5)
        add_code_block(doc, item["prompt"])
        add_bullets(doc, item["observe"])
        add_text(doc, item["safety"], bold_prefix=labels["safety_label"], after=8)


def add_tests(doc, labels):
    add_heading(doc, labels["tests"], 1)
    add_text(doc, labels["test_intro"], after=8)
    add_table(doc, labels["test_headers"], labels["test_rows"], [820, 2220, 2760, 1840, 1720])
    add_callout(doc, labels["pass_label"], labels["pass_criteria"])


def add_runbook(doc, labels):
    add_heading(doc, labels["runbook"], 1)
    add_heading(doc, labels["before"], 2)
    add_bullets(doc, labels["before_items"])
    add_heading(doc, labels["fallback"], 2)
    add_bullets(doc, labels["fallback_items"])
    add_heading(doc, labels["close"], 2)
    add_callout(doc, labels["closing_label"], labels["closing_quote"])


def english_labels():
    return {
        "architecture": "1. Product architecture to explain",
        "message": "Core message",
        "architecture_message": "MergePilot is a local-first desktop DevOps agent: it connects a selected local repository to its Azure DevOps context, gathers evidence through typed tools, and keeps every write action behind an approval boundary.",
        "layer": "Layer", "responsibility": "Responsibility",
        "desktop": "Tauri + React desktop", "desktop_desc": "Owns the Project Link, chat surface, transcript, approvals, and the operator experience.",
        "daemon": "Local daemon / SSE", "daemon_desc": "Receives requests locally, emits ordered Turn events, persists safe session state, and streams updates to the desktop.",
        "runtime": "Turn runtime + models", "runtime_desc": "A fast narrator produces short public action narratives; the main agent reasons over evidence, chooses tools, and produces the final answer.",
        "integrations": "Typed tools + MCP", "integrations_desc": "Local Git and repository context are combined with Azure DevOps and optional web/MCP reads. Actual calls, not predicted calls, appear in the transcript.",
        "policy": "Approval + local-first policy", "policy_desc": "Read-only work can proceed; commit, push, PR updates, and pipeline triggers remain explicit approval decisions. Credentials remain in local configuration or secret storage.",
        "flow_label": "Runtime flow: ", "flow": "Runtime flow: Developer -> Desktop -> local daemon/SSE -> Turn runtime -> Git, Azure DevOps MCP, optional web research -> evidence-backed final answer.",
        "turn": "2. The story of one visible Turn",
        "speaker_line": "What to say",
        "turn_quote": "We do not hide automation behind a chat answer. One user message becomes one recoverable Turn: public intent, actual evidence gathering, approval where needed, and a bounded conclusion.",
        "turn_steps": [
            "The sender sees Working for 0s immediately. This is local UI feedback, not a fake model response.",
            "The narrator streams a short, task-specific public action narrative. It is not private chain-of-thought and not a fixed system phrase.",
            "Only when a tool actually starts does a Ran commands group appear. Its commands are ordered and expandable for evidence.",
            "New evidence may produce a new public narrative and a new command group. The system never prints a future command list in advance.",
            "When execution seals, Working auto-collapses. The final conclusion streams outside it; copy and the end time appear only after the Turn finishes.",
        ],
        "guard_label": "Visible-data rule: ", "turn_guard": "Visible-data rule: The transcript contains public intent, actual actions, and bounded user-safe summaries. It excludes private reasoning, raw model payloads, secrets, and raw tool ledgers in the final answer.",
        "agenda": "3. 15-minute presentation run-of-show",
        "time": "Time", "segment": "Segment", "talk_track": "Talk track",
        "agenda_rows": [
            ("0:00-2:00", "The problem", "Local code, pull requests, policies, and pipelines are fragmented. Automation is useful only when it is scoped and reviewable."),
            ("2:00-5:00", "Product and architecture", "Introduce Project Link, local-first execution, the Turn transcript, MCP integrations, and explicit approval."),
            ("5:00-9:00", "Demo A: deployment-risk review", "Show real Git evidence, two-stage narration, command groups, collapse, and final conclusion."),
            ("9:00-12:00", "Demo B: PR readiness", "Show Azure DevOps evidence alongside local context without turning the chat into a log viewer."),
            ("12:00-14:00", "Demo C: approval boundary", "Prepare a pipeline trigger or equivalent write action; explain why nothing runs before approval."),
            ("14:00-15:00", "Close", "Reinforce auditability: a single Turn makes intent, evidence, policy, and outcome inspectable."),
        ],
        "rule": "Presentation rule", "rule_text": "Do not demo every feature. Use two read-only evidence cases and one approval case on a dedicated non-production Project Link.",
        "demo": "4. Live demonstration scripts",
        "goal_label": "Goal: ", "safety_label": "Safety boundary: ",
        "demos": [
            {"title": "Demo A - Review deployment risk from local Git evidence", "goal": "Goal: show a complete read-only Turn against a selected Project Link, not the MergePilot source repository.", "prompt": "Assess deployment risk in this Project Link's changed configuration.\nFirst establish the change scope, then inspect only the relevant diff.\nReport evidence and recommendations. Do not modify files or remote state.", "observe": ["Working appears immediately; the first public narrative is task-specific rather than a fixed loading label.", "One or more actual Git command groups appear only after commands begin.", "The final conclusion is outside the collapsed Working area and does not replay raw stdout or configuration snippets."], "safety": "Safety boundary: no stage, commit, push, or remote mutation is allowed."},
            {"title": "Demo B - Assess Azure DevOps pull request readiness", "goal": "Goal: show that Project Link binds local branch context to PR, policy, thread, and pipeline evidence.", "prompt": "Review the active pull request for release readiness.\nCheck unresolved review threads, required policies, and the latest pipeline result.\nKeep this read-only and identify only evidence-backed blockers.", "observe": ["Azure DevOps reads appear in the same Turn as local evidence, with connector-aware activity.", "A new narrative can appear after initial evidence before the next action group.", "The final answer states blockers and evidence without exposing an API payload."], "safety": "Safety boundary: no PR field, reviewer, label, policy, or pipeline state is modified."},
            {"title": "Demo C - Show the approval boundary", "goal": "Goal: make write safety visible rather than claiming it in a slide.", "prompt": "Prepare a request to trigger pipeline <demo-pipeline-id> for the current branch.\nShow the exact target and risk, but do not run it until I approve.", "observe": ["The agent gathers target and risk evidence first.", "An approval activity remains in the same Turn; the pipeline has not started yet.", "Only a deliberate approval may continue the action; the final record explains the outcome."], "safety": "Safety boundary: use a disposable demonstration pipeline, never a production deployment."},
        ],
        "tests": "5. High-quality demonstration test cases",
        "test_intro": "A good live test validates the observable lifecycle, not merely whether the final prose sounds plausible. Run these against prepared Project Links and record Pass, Partial, or Fail.",
        "test_headers": ["ID", "Precondition", "Prompt / action", "Expected visible lifecycle", "Must not happen"],
        "test_rows": [
            ("T-01", "Dirty demo repo", "Review the current working tree and identify release risks. Read-only only.", "Narrative -> actual Git group -> evidence-based conclusion", "Any Git write"),
            ("T-02", "PR with a policy/thread", "Assess the active pull request for merge readiness.", "Local + ADO evidence remains in one ordered Turn", "PR mutation"),
            ("T-03", "Failed demo pipeline", "Investigate the latest failed pipeline and identify the most likely failing stage.", "Pipeline read -> bounded diagnosis -> next safe action", "Raw full logs in final"),
            ("T-04", "Demo pipeline available", "Prepare a request to trigger pipeline <id>; do not run until I approve.", "Evidence -> approval -> optional continuation", "Trigger before approval"),
            ("T-05", "ADO disconnected or invalid", "Check pull request readiness and explain any blocker.", "Graceful error closes the Turn with recovery guidance", "Stuck Working or raw HTTP body"),
            ("T-06", "Long command / cancel enabled", "Run the approved validation command, then cancel it.", "Running state -> cancelled terminal state -> collapsed Working", "Duplicate final or stale timer"),
        ],
        "pass_label": "Universal pass criteria", "pass_criteria": "Working is visible immediately; each command is real and ordered; no future command list is pre-rendered; every terminal state closes the Turn; final text is separate from evidence; Copy and end time appear only after completion; secrets and raw payloads do not appear.",
        "runbook": "6. Operator runbook and fallback plan",
        "before": "Before the audience arrives", "before_items": ["Use a dedicated non-production repository and Project Link. Do not use MergePilot's own repository as the target.", "Pre-authenticate Azure DevOps and verify the selected PR, policy, and demo Pipeline IDs.", "Use deterministic data: 2-4 changed files, one visible PR condition, and one harmless write approval scenario.", "Validate the desktop sidecar, local daemon health, selected Project Link, and model configuration before the session.", "Prepare a recorded run of the exact same Project Link as a network/provider fallback."],
        "fallback": "If the network or model is slow", "fallback_items": ["Do not replace the delay with a fixed synthetic narrative. Explain that Working is immediate and public narration comes from the model.", "Use the pre-recorded run for the rest of the flow, then return to the live UI for Project Link and approval controls.", "Treat a slow or failed provider as a test result: show graceful terminal state and recovery guidance rather than hiding it."],
        "close": "Closing line", "closing_label": "Say this", "closing_quote": "MergePilot does not hide automation behind a chat answer. Every decision, evidence-gathering action, approval boundary, and final conclusion remains visible inside one recoverable Turn.",
    }


def chinese_labels():
    # Prompts stay English deliberately: the product's default working language is English.
    data = english_labels()
    data.update({
        "architecture": "1. 演讲应说明的产品架构",
        "message": "核心信息",
        "architecture_message": "MergePilot 是一个本地优先的桌面 DevOps Agent：它把选定的本地仓库与 Azure DevOps 上下文连接起来，通过类型化工具收集证据，并把所有写操作置于明确审批边界之后。",
        "layer": "层级", "responsibility": "职责",
        "desktop": "Tauri + React 桌面端", "desktop_desc": "承载 Project Link、对话界面、Turn Transcript、审批与操作者体验。",
        "daemon": "本机 daemon / SSE", "daemon_desc": "本地接收请求、发出有序 Turn 事件、持久化安全会话状态，并把流式更新发送到桌面端。",
        "runtime": "Turn Runtime + 模型", "runtime_desc": "快速 Narrator 输出简短公开行动叙述；主 Agent 基于证据进行判断、选工具并生成最终结论。",
        "integrations": "类型化工具 + MCP", "integrations_desc": "本地 Git 与仓库上下文可组合 Azure DevOps 及可选 Web/MCP 读取；Transcript 只显示实际调用，不显示预测调用。",
        "policy": "审批 + 本地优先策略", "policy_desc": "只读工作可直接执行；提交、推送、PR 更新和触发 Pipeline 都是显式审批决策。凭据仅存在本机配置或密钥存储。",
        "flow_label": "运行路径：", "flow": "运行路径：开发者 -> Desktop -> 本机 daemon/SSE -> Turn Runtime -> Git、Azure DevOps MCP、可选 Web Research -> 基于证据的最终结论。",
        "turn": "2. 一个可见 Turn 的故事",
        "speaker_line": "可直接这样讲",
        "turn_quote": "我们不会把自动化藏在一句聊天答案后。一个用户消息会成为一个可恢复的 Turn：公开意图、真实证据收集、必要时的审批，以及受边界控制的结论。",
        "turn_steps": [
            "发送后立即可见 Working for 0s。这是本地 UI 反馈，不是假冒的模型回复。",
            "Narrator 流式输出短小、针对任务的公开行动叙述；它不是私有推理链，也不是固定系统文案。",
            "只有工具真正开始后才显示 Ran commands；命令严格保序，并可展开查看证据。",
            "新证据可以触发新的公开叙述与新的命令组。系统绝不会预先打印未来命令清单。",
            "执行期封口后，Working 自动收起；最终结论在外部流式输出，只有 Turn 完成才出现复制图标与结束时间。",
        ],
        "guard_label": "可见数据规则：", "turn_guard": "可见数据规则：Transcript 只包含公开意图、真实行动和有边界的用户安全摘要；它排除私有推理、原始模型 payload、密钥，以及最终结论中的原始工具账本。",
        "agenda": "3. 15 分钟演讲流程",
        "time": "时间", "segment": "环节", "talk_track": "讲述要点",
        "agenda_rows": [
            ("0:00-2:00", "问题", "本地代码、PR、策略与流水线分散；自动化只有在可控、可审查时才真正有价值。"),
            ("2:00-5:00", "产品与架构", "介绍 Project Link、本地优先执行、Turn Transcript、MCP 集成与显式审批。"),
            ("5:00-9:00", "演示 A：部署风险审查", "展示真实 Git 证据、分阶段叙述、命令组、自动收起与最终结论。"),
            ("9:00-12:00", "演示 B：PR 就绪度", "展示 Azure DevOps 证据与本地上下文如何融合，而不是把对话变成日志查看器。"),
            ("12:00-14:00", "演示 C：审批边界", "准备触发 Pipeline 或等价写操作，说明为何审批前绝不会执行。"),
            ("14:00-15:00", "收尾", "强调可审计性：单一 Turn 令意图、证据、策略和结果都可追溯。"),
        ],
        "rule": "演示规则", "rule_text": "不要现场覆盖全部功能。使用一个专用非生产 Project Link，重点展示两条只读证据链和一条审批链。",
        "demo": "4. 现场演示脚本",
        "goal_label": "目标：", "safety_label": "安全边界：",
        "demos": [
            {"title": "演示 A - 基于本地 Git 证据审查部署风险", "goal": "目标：在选定的 Project Link 上展示完整只读 Turn，而不是审查 MergePilot 源码仓库。", "prompt": "Assess deployment risk in this Project Link's changed configuration.\nFirst establish the change scope, then inspect only the relevant diff.\nReport evidence and recommendations. Do not modify files or remote state.", "observe": ["Working 立即出现；首段公开叙述与任务有关，而不是固定 loading 标签。", "仅在命令启动后才出现一个或多个真实 Git 命令组。", "最终结论位于已收起的 Working 区外，且不复述原始 stdout 或配置片段。"], "safety": "安全边界：禁止 stage、commit、push 或任何远端变更。"},
            {"title": "演示 B - 评估 Azure DevOps PR 就绪度", "goal": "目标：展示 Project Link 如何绑定本地分支上下文与 PR、策略、评论和流水线证据。", "prompt": "Review the active pull request for release readiness.\nCheck unresolved review threads, required policies, and the latest pipeline result.\nKeep this read-only and identify only evidence-backed blockers.", "observe": ["Azure DevOps 读取与本地证据处于同一个 Turn，并保留 connector 信息。", "初始证据之后、下一行动组之前，可以出现新的公开叙述。", "最终答案只报告阻塞项和证据，不暴露 API payload。"], "safety": "安全边界：不修改 PR 字段、Reviewer、标签、策略或 Pipeline 状态。"},
            {"title": "演示 C - 展示审批边界", "goal": "目标：让写操作安全性在真实界面中可见，而不仅是幻灯片中的宣称。", "prompt": "Prepare a request to trigger pipeline <demo-pipeline-id> for the current branch.\nShow the exact target and risk, but do not run it until I approve.", "observe": ["Agent 先收集目标与风险证据。", "审批活动仍属于同一个 Turn，Pipeline 在此时尚未启动。", "只有明确批准才能继续；最终记录会说明结果。"], "safety": "安全边界：使用可弃用的演示 Pipeline，绝不用生产部署。"},
        ],
        "tests": "5. 高质量现场测试用例",
        "test_intro": "高质量现场测试验证可观察的生命周期，而不只是最终文字是否听起来合理。请在准备好的 Project Link 上执行，并记录 Pass、Partial 或 Fail。",
        "test_headers": ["ID", "前置条件", "Prompt / 操作", "预期可见生命周期", "禁止发生"],
        "test_rows": [
            ("T-01", "有改动的演示仓库", "Review the current working tree and identify release risks. Read-only only.", "叙述 -> 真实 Git 组 -> 基于证据的结论", "任何 Git 写操作"),
            ("T-02", "含策略/评论的 PR", "Assess the active pull request for merge readiness.", "本地 + ADO 证据处于同一有序 Turn", "修改 PR"),
            ("T-03", "失败的演示 Pipeline", "Investigate the latest failed pipeline and identify the most likely failing stage.", "读取 Pipeline -> 有边界诊断 -> 下一安全动作", "原始完整日志进入 final"),
            ("T-04", "可用演示 Pipeline", "Prepare a request to trigger pipeline <id>; do not run until I approve.", "证据 -> 审批 -> 可选继续执行", "审批前触发"),
            ("T-05", "ADO 断连或无效", "Check pull request readiness and explain any blocker.", "优雅错误关闭 Turn 并给出恢复建议", "Working 卡死/原始 HTTP body"),
            ("T-06", "长命令/可取消", "Run the approved validation command, then cancel it.", "运行态 -> 取消终态 -> 收起 Working", "重复 final/旧计时器"),
        ],
        "pass_label": "通用通过标准", "pass_criteria": "Working 立即可见；每条命令都真实发生且有序；不预渲染未来命令；每个终态都关闭 Turn；最终文本与证据分离；Copy 和结束时间只在完成后出现；不展示密钥与原始 payload。",
        "runbook": "6. 操作者清单与故障预案",
        "before": "观众到场前", "before_items": ["使用专用非生产仓库和 Project Link，不要将 MergePilot 自身仓库作为目标。", "预先完成 Azure DevOps 登录，并核实所选 PR、策略和演示 Pipeline ID。", "准备确定性数据：2-4 个改动文件、一个可见 PR 条件和一个无害写审批场景。", "在会前验证 Desktop sidecar、本机 daemon 健康、所选 Project Link 和模型配置。", "录制同一 Project Link 的完整成功流程，作为网络/模型提供商异常时的备份。"],
        "fallback": "网络或模型变慢时", "fallback_items": ["不要用固定假叙述掩盖延迟；说明 Working 是即时的，而公开叙述来自真实模型。", "使用预录流程完成剩余演示，再返回实时 UI 展示 Project Link 与审批控件。", "把慢速或失败的 provider 视为测试结果：展示优雅终态与恢复建议，而不是隐藏异常。"],
        "close": "收尾", "closing_label": "可直接这样说", "closing_quote": "MergePilot 不会把自动化藏在一条聊天答案后。每个决策、证据收集动作、审批边界与最终结论都可见，并且处于同一个可恢复的 Turn 中。",
    })
    return data


def build(path, labels, title, subtitle, stripe, header_text, footer_text):
    doc = setup_doc(header_text, footer_text)
    add_title(doc, title, subtitle, stripe)
    add_architecture(doc, labels)
    add_turn(doc, labels)
    add_agenda(doc, labels)
    doc.add_page_break()
    add_demo(doc, labels)
    add_tests(doc, labels)
    doc.add_page_break()
    add_runbook(doc, labels)
    doc.save(path)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    build(
        OUT / "MergePilot_Live_Demo_Playbook_EN.docx",
        english_labels(),
        "MergePilot Live Demo Playbook",
        "A 15-minute presentation, architecture narrative, and evidence-based demonstration guide",
        [("FORMAT", "15-minute live demo"), ("AUDIENCE", "Engineering and delivery leaders"), ("OPERATING MODE", "Non-production Project Link")],
        "MergePilot | Live Demo Playbook",
        "Internal demonstration guide | Use a non-production Project Link",
    )
    build(
        OUT / "MergePilot_现场演示指南_CN.docx",
        chinese_labels(),
        "MergePilot 现场演示指南",
        "15 分钟演讲结构、架构叙述与基于证据的现场演示手册",
        [("形式", "15 分钟现场演示"), ("受众", "工程与交付负责人"), ("运行模式", "非生产 Project Link")],
        "MergePilot | 现场演示指南",
        "内部演示手册 | 请使用非生产 Project Link",
    )


if __name__ == "__main__":
    main()
